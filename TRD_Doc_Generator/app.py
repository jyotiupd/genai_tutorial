from flask import Flask, render_template, request, send_from_directory, jsonify, send_file
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
from reportlab.pdfgen import canvas
import os
import io
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

 

app = Flask(__name__)

 

# Replace with your OpenAI API key
#openai.api_key = ''
os.environ['OPENAI_API_KEY'] = ''
openai.api_key = os.getenv('OPENAI_API_KEY')
DOC_DIR = "./generated_docs"
@app.route('/')
def index():
    return open('templates/index.html').read()

 

 

@app.route('/process', methods=['POST'])
def process():
    if request.method == 'POST':
        query = request.form.get('data')
        print(query)
        p1="""
            - Introduction: An overview of the purpose and scope of the document.
            - System Overview: A brief description of the system or product."""
        p2="""
            - Functional Requirements: Describes what the system or product should do. This might include user stories or use cases.
            - Non-Functional Requirements: These are the quality attributes the system must have, such as performance, security, reliability, etc."""
        p3="""
            - Interface Requirements: Details of how the system will interact with other systems, if applicable.
            - Performance Requirements: Specific performance benchmarks the system must meet."""
        p4="""
            - Data Requirements: Describes the type of data the system will handle and how.
            - Hardware and Software Requirements: The minimum or required hardware and software configurations for the system to operate."""
        p5="""
            - Security Requirements: Details on how the system will ensure data protection and user privacy.
            - Constraints and Assumptions: Any limitations or assumptions made during the drafting of the requirements."""
        
        prompt = prompt = PromptTemplate(input_variables = ['query','points'], template =

                          """Provide the Technical Requirements for {query} in detail with the help of given instruction \
                            for points delimited by triple backticks,
                            instructions :
                            1. generate the output only for the points provided
                            2. Donot add any extra details apart from the details for the mentioned points
                            3. Donot add any extra section
                            4. Donot repeate the prompt in the output
                            5. Generate the output in text format

                          ```{points}```
                          """)
        
        llm = OpenAI(temperature = .9)
        #How can I access GPT-4?

        #https://help.openai.com/en/articles/7102672-how-can-i-access-gpt-4
        chain = LLMChain(llm = llm, prompt = prompt,verbose = True)
        #response  = chain.run(query = query)

        chain = LLMChain(llm = llm, prompt = prompt,verbose = True)
        res1  = chain.run({'query' : query,'points':p1})
        res2  = chain.run({'query' : query,'points':p2})
        res3  = chain.run({'query' : query,'points':p3})
        res4  = chain.run({'query' : query,'points':p4})
        res5  = chain.run({'query' : query,'points':p5})
        heading = "Technical Requirements for {query}".format(query = query)
        response = "\n".join([heading,res1,res2,res3,res4,res5])
        return jsonify({"data": response})
        #doc = Document()
        #doc.add_heading('Technical Requirements Document', 0)
        #doc.add_paragraph(response)
        #filename = "response.docx"
        #doc_path = os.path.join(DOC_DIR, filename)

        #doc.save(doc_path)
        #return send_from_directory(directory=DOC_DIR, path=filename, as_attachment=True)

 

    #return render_template('index.html')

 

 

@app.route('/download/<filetype>', methods=['POST'])
def download(filetype):
    data = request.form.get('data')

    if filetype == 'pdf':

        output = io.BytesIO()

        doc = SimpleDocTemplate(output, pagesize=letter)

 

        # Define styles

        styles = getSampleStyleSheet()

        headline_style = styles['Heading1']

        normal_style = styles['BodyText']

 

        # Split the input data into lines

        lines = data.split('\n')

 

        # Convert lines to Flowable objects

        flowables = []

        for line in lines:

            if line.startswith('*') and line.endswith('*'):  # assuming this line is a headline

                flowables.append(Paragraph(line[1:-1], headline_style))

                flowables.append(Spacer(1, 12))

            else:

                flowables.append(Paragraph(line, normal_style))

                flowables.append(Spacer(1, 12))

 

        # Build the PDF

        doc.build(flowables)

 

        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'response.{filetype}', mimetype='application/pdf')

    elif filetype == 'doc':
        # Using python-docx to generate a DOCX
        doc = Document()
        doc.add_paragraph(data)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f'response.{filetype}', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

 

 

if __name__ == '__main__':
    if not os.path.exists(DOC_DIR):
        os.makedirs(DOC_DIR)
    app.run(debug=True)